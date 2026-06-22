"""
Парсер файлов: PDF, DOCX, XLSX/XLS/CSV, TXT, MD и др.
"""

import io
from pathlib import Path


def parse_file(filename: str, content_bytes: bytes) -> str:
    """Преобразует бинарный файл в plain text."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(content_bytes)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(content_bytes)
    elif suffix in (".xlsx", ".xls"):
        return _parse_excel(content_bytes)
    elif suffix == ".csv":
        return content_bytes.decode("utf-8", errors="replace")
    else:
        return content_bytes.decode("utf-8", errors="replace")


def _parse_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Установите pdfplumber: pip install pdfplumber")

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = []
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
    if not pages:
        raise ValueError("PDF не содержит извлекаемого текста (возможно, отсканирован)")
    return "\n\n".join(pages)


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Установите python-docx: pip install python-docx")

    doc = Document(io.BytesIO(data))
    parts = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def _parse_excel(data: bytes) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("Установите openpyxl: pip install openpyxl")

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(v) for v in row if v is not None and str(v).strip())
            if row_text.strip():
                rows.append(row_text)
        if rows:
            sheets.append(f"=== {sheet_name} ===\n" + "\n".join(rows))

    wb.close()
    return "\n\n".join(sheets)
